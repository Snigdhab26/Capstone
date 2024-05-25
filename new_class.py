import streamlit
from fpdf import FPDF
from datetime import datetime
import pandas as pd
import re
import os
from docx import Document
import PyPDF2
import requests
from bs4 import BeautifulSoup
from textblob import TextBlob
from transformers import AutoTokenizer, AutoModel
from transformers import pipeline
import torch
from pymed import PubMed
import dateparser
from tabula import read_pdf
import tabula as t
from PyPDF2 import PdfFileReader
from io import BytesIO
import pandas as pd
from nltk.translate.bleu_score import sentence_bleu, SmoothingFunction
from rouge_score import rouge_scorer
from bert_score import score as bert_score
from sacrebleu import BLEU
from transformers import pipeline
from nltk.translate import meteor_score as nltk_meteor_score
import nltk

PRRER_No = 'PBRER No. {}'.format('P3')
today_date = datetime.today()

class PDF(FPDF):
    def __init__(self, product_name, df1 , df2,formatted_begin_date, formatted_end_date):
        super().__init__('P', 'mm', 'Letter')
        self.product_name = product_name
        self.formatted_date = today_date.strftime("%d-%b-%Y")
        self.detail = df1
        self.formatted_begin_date = formatted_begin_date
        self.formatted_end_date = formatted_end_date
    def header(self):
        # Left-aligned lines
        self.set_font('helvetica', '', 8)
        self.set_text_color(0, 0, 0)
        self.cell(0, 1, 'Sandoz', ln=2, align='L')
        self.cell(0, 1, 'Reporting period {} to {}'.format(self.formatted_begin_date, self.formatted_end_date), ln=2, align='R')
        self.ln(1)
        self.cell(0, 1, self.product_name, ln=2, align='L')
        self.cell(0, 1, 'Date: {}'.format(self.formatted_date), ln=2, align='R')
        self.ln(2)
        self.cell(0, 1, PRRER_No, align='L')
        self.cell(0, 1, f'Page {self.page_no()}', align='R')
        self.ln(5)
        self.line(self.l_margin, self.t_margin + 10, self.w - self.r_margin, self.t_margin + 10)
        self.ln(5)

    # Page footer
    def footer(self):
        # Set position of the footer
        self.set_y(-15)
        self.set_font('helvetica', '', 8)
        self.set_text_color(0, 0, 0)
        # Page number
        self.cell(0, 10,'CONFIDENTIAL', align='C')

    def chapter_title(self, title, link):
        self.set_link(link)
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, title.encode('latin-1',  'replace').decode('latin-1'), 0, 1, 'L')  # Encode and decode using latin-1
        self.ln(4)

    def chapter_body(self, body,link):
        self.set_link(link)
        self.set_font('Arial', '', 12)
        self.multi_cell(0, 10, body.encode('latin-1', 'replace').decode('latin-1'))  # Encode and decode using latin-1
        self.ln()


    def filename_extract(self, word, filename):
        # Create the regular expression pattern to match the word as a whole word
        pattern = re.compile(re.escape(word.upper()), re.IGNORECASE)
        # print(pattern)
        # Filter filenames based on the pattern
        matching_files = [f for f in filename if re.search(pattern, f.upper())]

        # Return the first matching filename or None if no match found
        return matching_files[0] if matching_files else None

    def Index_page(self, title, p,margin, link ,color ='black'):
        self.set_text_color(*color)
        self.set_font('Arial', 'B', 5)
        self.set_text_color(0, 0, 255)
        self.cell(0, 8, title.encode('latin-1', 'replace').decode('latin-1'), 0, 1, p,margin,link = link)  # Encode and decode using latin-1
        self.set_text_color(0, 0, 0)

    # def handle_encoding(self,cell):
    #     try:
    #        # Try encoding and decoding with 'latin-1' encoding
    #        encoded_cell = cell.encode('latin-1', errors='ignore')
    #        decoded_cell = encoded_cell.decode('latin-1', errors='ignore')
    #        return decoded_cell
    #     except UnicodeEncodeError:
    #     # If encoding error occurs, return original cell value
    #        return cell

    def add_table(self, df):
        # Calculate the maximum length of the values in the DataFrame
        max_length = 0

        # Calculate maximum length manually
        for _, row in df.iterrows():
            for cell in row:
                #print(cell)
                cell_length = len(str(cell))
                #print(cell_length)
                if cell_length > max_length:
                    max_length = cell_length

        self.set_font('helvetica', '', 8)
        self.set_text_color(0, 0, 0)
        for i, (_, row) in enumerate(df.iterrows()):
            for j, cell in enumerate(row):
                #print(j,cell)
                # Adjust cell width based on index
                cell_width = max_length/2 if j == 0 else max_length * 1.5
                # Use calculated cell width and fixed height
                self.cell(cell_width, 10, str(cell).encode('latin-1', errors='ignore').decode('latin-1') , border=1)
            self.ln()



    def write_dataframe_to_pdf(self, df, output_path):
        data = [df.columns.tolist()] + df.values.tolist()
        table = pd.DataFrame(data)  # Convert data to DataFrame

        # Call add_table method to add DataFrame to PDF
        self.add_table(table)

    def extract_text_from_page(self,pdf_file, page_number):
                with open(pdf_file, 'rb') as file:
                    #print('1')
                    pdf_reader = PyPDF2.PdfReader(file)
                    page = pdf_reader.pages[page_number]
                    return page.extract_text().encode('latin-1', errors='ignore').decode('latin-1')

    
    def Abbreviations(self, con, output_path, filename):
        abbreviation = con
        print(filename)
        filtered_detail = 'RMP'
        print(filtered_detail)
        filtered_filenames = [f for f in filename if
                              filtered_detail in f]  # [f for f in filename if any(filename in f for filename in filtered_detail)]
        # Define the PDF file path
        pdf_file_path = None
        for f in filtered_filenames:
            pdf_file_path = f
        pages_with_abbreviation = []
        print(pdf_file_path)
        # Extract text from each page and check for the abbreviation
        try:
            # Get the total number of pages in the PDF using PyPDF2
            with open(pdf_file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
            print(total_pages)
            # Loop through each page to extract text and check for the abbreviation
            for page_number in range(total_pages):
                text = self.extract_text_from_page(pdf_file_path, page_number)
                # print(text)
                if abbreviation.upper() in text.upper():
                    if "Table of contents".upper() not in text.upper():
                        # print(abbreviation.upper())
                        # print(text.upper())
                        pages_with_abbreviation.append(
                            page_number + 1)  # Add 1 to convert zero-based index to page number
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
        if len(pages_with_abbreviation) == 1 or len(pages_with_abbreviation) == 0:
            pagenums = pages_with_abbreviation if len(pages_with_abbreviation) == 1 else None
        else:
            pagenums = max(pages_with_abbreviation) - min(pages_with_abbreviation)

        if pagenums != None:
            text = t.read_pdf(pdf_file_path, pages=pagenums, multiple_tables=True)
            table = pd.DataFrame()
            table = pd.concat([table, text[0]], ignore_index=False)
            df1 = pd.DataFrame(table)
            self.write_dataframe_to_pdf(df1, output_path)
            # return df1
        else:
            df1 = pd.DataFrame()
        return df1
    # def filename_extract(self,word, filename):
    #         # Create the regular expression pattern to match the word as a whole word
    #         print(word)
    #         print(filename)
    #         pattern = re.compile(re.escape(word.upper()), re.IGNORECASE)
    #         print(pattern)
    #         # Filter filenames based on the pattern
    #         matching_files = [f for f in filename if re.search(pattern, f.upper())]
    #
    #         # Return the first matching filename or None if no match found
    #         return matching_files[0] if matching_files else None

    def convert_table_to_html(self, docx_file):
        try:
            document = Document(docx_file)
            docx_content = ''
            # docx_content = "\n".join([paragraph.text for paragraph in document.paragraphs])
            # Check if the document has any tables
            if len(document.tables) > 0:
                # Initialize the HTML string with the table structure

                html_string = '<table border="1">\n'

                # Iterate through each table in the document
                for table in document.tables:
                    # Process nested tables within the current table
                    for row in table.rows:
                        for cell in row.cells:
                            # Check for nested tables within the cell
                            html_string += f'<td>{cell.text}</td>\n'

                            for nested_table in cell.tables:
                                for nested_row in nested_table.rows:
                                    html_string += '<tr>\n'

                                    # Iterate through each cell in the row
                                    for nested_cell in nested_row.cells:
                                        html_string += f'<td>{nested_cell.text}</td>\n'

                                    html_string += '</tr>\n'

                    html_string += '</tr>\n'

                # Close the HTML table structure
                html_string += '</table>'
            else:
                html_string = ''

            docx_content += (html_string)
            return docx_content
        except Exception as e:
            print(f"An error occurred: {e}")
            return None

    def Section2(self, con, output_path, filename, drug):
        try:
            excel_filepath = self.filename_extract(con, filename)
            excel_data = pd.read_excel(excel_filepath)
            excel_data['Approval Date'] = pd.to_datetime(excel_data['Approval Date'], format='%d.%m.%Y',
                                                         errors='coerce')

            # Drop rows with missing or incorrect date formats (if any)
            excel_data.dropna(subset=['Approval Date'], inplace=True)

            # Sort DataFrame by 'Approval Date' column in ascending order
            excel_data_sorted = excel_data.sort_values(by='Approval Date')

            # Get the first row (earliest approval date)
            earliest_approval_row = excel_data_sorted.iloc[0]

            # Extract the 'Country Name' and 'Approval Date' from the earliest approval row
            country_with_earliest_approval = earliest_approval_row['Country Name']
            earliest_approval_date = earliest_approval_row['Approval Date']

            # Format the date as "DD Month YYYY" (e.g., "02 March 2017")
            formatted_approval_date = earliest_approval_date.strftime('%d %B %Y')

            print("Country with earliest Approval Date:", country_with_earliest_approval)
            print("Earliest Approval Date:", formatted_approval_date)

            # Filter DataFrame where 'Marketing Status Text' is 'Marketed'
            marketed_df = excel_data[excel_data['Marketing Status Text'] == 'Marketed']

            # Count unique 'Country Name' where 'Marketing Status Text' is 'Marketed'
            marketed_count = marketed_df['Country Name'].nunique()

            # Filter DataFrame where 'Registration Status Text' is 'Approved'
            approved_df = excel_data[excel_data['Registration Status Text'] == 'Approved']

            # Count unique 'Country Name' where 'Registration Status Text' is 'Approved'
            approved_count = approved_df['Country Name'].nunique()

            print("Unique 'Country Name' where 'Marketing Status Text' is 'Marketed':", marketed_count)
            print("Unique 'Country Name' where 'Registration Status Text' is 'Approved':", approved_count)

            # Save the counts in variables as requested
            marketed = marketed_count
            approved = approved_count

            # Prepare the query text
            query_text = f"\n{drug} was first registered in {country_with_earliest_approval} on {formatted_approval_date}. {drug} is currently marketed by Sandoz in {marketed_count} countries and has approval for market in {approved_count} countries.\n\nThe registered indications of {drug} are presented in Section 1. Details about recommended posology are described in Appendix 1.\n\nThe worldwide market authorisation status of {drug} is presented in Appendix 2."

            # Add the query text to the PDF
            self.set_font('Arial', '', 12)
            self.set_text_color(0, 0, 0)
            self.multi_cell(0, 10, query_text, 0, 1)  # Encode and decode using latin-1
            self.set_text_color(0, 0, 0)  # Reset text color
            return marketed, approved, query_text
        except FileNotFoundError:
            print("The specified file does not exist.")
        except Exception as e:
            print("An error occurred:", e)


    def Section3(self, con, content, output_path, filename, drug):
        try:
            _filepath = self.filename_extract(con, filename)
            print(_filepath)
            _, file_extension = os.path.splitext(_filepath)
            print(file_extension)

            if file_extension.lower() == '.docx':
                doc = Document(_filepath)
                total_pages = len(doc.element.body.xpath('//w:sectPr'))
                content_words = content.split()
                print(content_words)
                # Iterate through each word in content
                tex = ''
                for i in range(len(content_words) % 2):  # Check for 3-4 word matches
                    pattern = ' '.join(content_words[i:i + 4])  # Create a pattern of 3-4 words
                    print(pattern)
                    # Iterate through each paragraph in the document
                    for paragraph in doc.paragraphs:
                        print(paragraph.text)
                        # Check if the pattern is found in the paragraph (case-insensitive)
                        if not re.search(re.escape(pattern), paragraph.text, flags=re.IGNORECASE):
                            # Skip this paragraph as it contains the pattern
                            if 'table:' in paragraph.text.lower():
                                print("1")
                                tex1 = pdf.convert_table_to_html(_filepath)
                                # print(tex1)
                                if len(tex1) > 0:
                                    clean_text = re.sub(r'<[^>]+>', '', tex1)
                                    clean_text = clean_text.replace('\n', '.')
                                    summary = summarizer(clean_text, max_length=75, min_length=50, length_penalty=2.0,
                                                         num_beams=4, no_repeat_ngram_size=2, early_stopping=True)
                                    result = summary[0]['summary_text']
                                    tex += result + '\n'
                            else:
                                print("2")
                                tex += paragraph.text + '\n'
                        else:
                            # Append the paragraph text to the filtered text
                            continue

                self.set_font('Arial', '', 12)
                self.set_text_color(0, 0, 0)
                self.multi_cell(0, 10, tex.encode('latin-1', errors='ignore').decode('latin-1'), 0,
                                1)  # Encode and decode using latin-1
                self.set_text_color(0, 0, 0)  # Reset text color
                return  tex
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")

    def Section4(self, con, content, output_path, filename, drug):
        try:
            _filepath = self.filename_extract(con, filename)
            print(_filepath)
            _, file_extension = os.path.splitext(_filepath)
            print(file_extension)

            if file_extension.lower() == '.docx':
                doc = Document(_filepath)
                total_pages = len(doc.element.body.xpath('//w:sectPr'))
                content_words = content.split()
                print(content_words)
                # Iterate through each word in content
                tex = ''
                result = ''
                clean_text = ''
                for i in range(len(content_words) % 2):  # Check for 3-4 word matches
                    pattern = ' '.join(content_words[i:i + 4])  # Create a pattern of 3-4 words
                    print(pattern)
                    # Iterate through each paragraph in the document
                    for paragraph in doc.paragraphs:
                        print(paragraph.text)
                        # Check if the pattern is found in the paragraph (case-insensitive)
                        if not re.search(re.escape(pattern), paragraph.text, flags=re.IGNORECASE):
                            # Skip this paragraph as it contains the pattern
                            if 'table:' in paragraph.text.lower():
                                print("1")
                                tex1 = pdf.convert_table_to_html(_filepath)
                                # print(tex1)
                                if len(tex1) > 0:
                                    clean_text = re.sub(r'<[^>]+>', '', tex1)
                                    clean_text = clean_text.replace('\n', '.')
                                    summary = summarizer(clean_text, max_length=75, min_length=50, length_penalty=2.0,
                                                         num_beams=4, no_repeat_ngram_size=2, early_stopping=True)
                                    result = summary[0]['summary_text']
                                    tex += result + '\n'
                            else:
                                print("2")
                                tex += paragraph.text + '\n'
                        else:
                            continue
                # print(tex)
                # print(tex)
                self.set_font('Arial', '', 12)
                self.set_text_color(0, 0, 0)
                self.multi_cell(0, 10, tex.encode('latin-1', errors='ignore').decode('latin-1'), 0,
                                1)  # Encode and decode using latin-1
                self.set_text_color(0, 0, 0)  # Reset text color
        except Exception as e:
            tex += clean_text + '\n'
            print(f"Error extracting text from PDF: {e}")
            self.set_font('Arial', '', 12)
            self.set_text_color(0, 0, 0)
            self.multi_cell(0, 10, tex.encode('latin-1', errors='ignore').decode('latin-1'), 0,
                            1)  # Encode and decode using latin-1
            self.set_text_color(0, 0, 0)  # Reset text color
        return tex

    def Section1(self, con, content, output_path, filename, drug):
        try:
            print(filename)
            print(drug)
            _filepath = self.filename_extract(con, filename)
            print(_filepath)
            _, file_extension = os.path.splitext(_filepath)
            # print(file_extension)
            date1 = self.formatted_begin_date
            date2 = self.formatted_end_date
            final_section1_text = f'This Periodic Safety Update Report (PSUR) on {drug} containing products was compiled for regulatory authorities in the Periodic Benefit-Risk Evaluation Report (PBRER) format detailed in the European Union (EU) and the International Council on Harmonisation (ICH)-E2C guidelines (Good Pharmacovigilance Practice Guideline Module VII Periodic Safety Update Reports, 2012 and ICH-E2C(R2). It summarizes the safety data received and processed by Sandoz from worldwide sources for the period covering {date1} to {date2}. The product is referred to as {drug} throughout this report.'
            doc = Document(_filepath)
            extracted_text = ""
            extracting = False
            extracting1 = False
            extracted_text2 = ''
            extracted_text3 = ''
            # Iterate through paragraphs
            for paragraph in doc.paragraphs:
                # Check if the paragraph contains the desired section identifiers
                if "Pharmacodynamic" in paragraph.text and "properties" in paragraph.text:
                    extracting = True
                elif "Pharmacokinetic " in paragraph.text and "properties" in paragraph.text:
                    extracting = False
                    break  # Exit the loop after finding the end of the section
                # If extracting is True, add the paragraph's text to the extracted text
                if extracting:
                    extracted_text += paragraph.text + "\n"
                    # print(extracted_text)
            if "Mechanism" in extracted_text:
                mechanism_index = extracted_text.find("Mechanism")
                # Find the index of ':' after "Mechanism of Action" header
                colon_index = extracted_text.find(":", mechanism_index)
                # Extract text from mechanism_index until colon_index
                text_beforemechanism_text = extracted_text[:mechanism_index].strip()
                col_index1 = extracted_text.find(":")
                extracted_text2 = extracted_text[colon_index + 1:].strip()
                colon_index = extracted_text2.find(":", mechanism_index)
                extracted_text3 = extracted_text2[:colon_index + 1].strip()
                extracting1 = True
            if extracting1:
                # print(1)
                final_section1_text += '\n' + text_beforemechanism_text + '\n' + extracted_text3 + '\n'
            else:
                # print(2)
                final_section1_text += '\n' + extracted_text + '\n'

            # print(final_section1_text)
            extracted_text = ""
            extracting = False
            # Iterate through paragraphs
            for paragraph in doc.paragraphs:
                # Check if the paragraph contains the desired section identifiers
                if "Therapeutic" in paragraph.text and "indications" in paragraph.text:
                    extracting = True
                elif "Posology" in paragraph.text and "method of administration" in paragraph.text:
                    extracting = False
                    # Include the paragraph containing "4.2" in the extracted text
                    # extracted_text += paragraph.text + "\n"
                    break  # Exit the loop after finding the end of the section
                # If extracting is True, add the paragraph's text to the extracted text
                if extracting:
                    extracted_text += paragraph.text + "\n"
            final_section1_text = final_section1_text + '\n' + extracted_text + '\n'
            # Print or do whatever you want with the extracted text
            # print(final_section1_text)
            paragraphs = final_section1_text.split('\n')
            filtered_paragraphs = [paragraph for paragraph in paragraphs if
                                   "5.1" not in paragraph and ":" not in paragraph and "4.1" not in paragraph]
            final_section1_text = '\n'.join(filtered_paragraphs)
            # print(final_section1_text)

            filtered_filenames = [f for f in filename if "RMP" in f]
            pdf_file_path = None
            if filtered_filenames:
                # Select the first filename from the list
                pdf_file_path = filtered_filenames[0]
                print(pdf_file_path)
            with open(pdf_file_path, 'rb') as file:
                # Create a PDF reader object
                reader = PyPDF2.PdfReader(file)
                page = reader.pages[0]
                page_text = page.extract_text()
                similar_text = ""
                # Define regex pattern to match text before drug name and after "Risk Management"
                pattern_before_drug = re.compile(r'.*{}\s*$'.format(re.escape(drug)), re.IGNORECASE | re.DOTALL)
                pattern_after_risk_management = re.compile(r'Risk\s+Management(?:[\s\S]*|$)', re.IGNORECASE)
                pattern_ = re.compile(r'Medical\s+Patient(?:[\s\S]*|$)', re.IGNORECASE)
                pattern_after_eu_safety = re.compile(r'EU\s+Safety(?:[\s\S]*|$)', re.IGNORECASE)
                text_after_drug = re.sub(pattern_before_drug, '', page_text)
                text_final = re.sub(pattern_after_risk_management, '', text_after_drug)
                text_final = re.sub(pattern_after_eu_safety, '', text_final)
                text_final = text_final.replace("Chief Medical Office & Patient Safety", "")
                summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
                summary = summarizer(text_final, max_length=75, min_length=50, length_penalty=2.0,num_beams=4, no_repeat_ngram_size=2, early_stopping=True)
                text_final = summary[0]['summary_text']

            final_section1_text += '\n' + text_final.strip() + '\n'
            final_section1_text += '\n' + 'Further details on mechanism of action, indications, special population, pharmaceutical forms and instructions and use in special populations are presented in the Appendix 1.' + '\n' + ' A PSUR on the same ingredient may be submitted by other Marketing Authorization Holders (MAHs). It needs to be taken into account that literature cases and case reports referring to unidentified generic products contained in this PSUR are likely to be the same as referenced by other MAHs.'
            self.set_font('Arial', '', 12)
            self.set_text_color(0, 0, 0)
            self.multi_cell(0, 10, final_section1_text.encode('latin-1', errors='ignore').decode('latin-1'), 0,
                            1)  # Encode and decode using latin-1
            self.set_text_color(0, 0, 0)  # Reset text color
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
        # Create a PDF reader object
        return final_section1_text

        # def Section6(con, output_path, filename, drug):

    def Section6(self, con, output_path, filename, drug):
        query_text = f"\n All {drug} cases were migrated into the new Sandoz Global Safety database in Jul-2023 as per of Sandoz Pharmacovigilance Separation from Novartis. \n This activity included consolidation of the Company product dictionaries, detection of duplicate cases and re-mapping of case report types. Such activities, however, may result in changes to the cumulative counts of data presented in previous aggregated safety reports including this PBRER. \n This report covering the interval period of {self.formatted_begin_date} to {self.formatted_end_date} represents a harmonization of methodologies to a single, standard Sandoz process which presents all information pertinent to an accurate assessment of the benefit/risk profile of allopurinol in a manner consistent with all Sandoz periodic aggregate reports. \n \n "
        # Add the query text to the PDF
        self.set_font('Arial', '', 12)
        self.set_text_color(0, 0, 0)
        self.multi_cell(0, 10, query_text.encode('latin-1', errors='ignore').decode('latin-1'), 0,
                        1)  # Encode and decode using latin-1
        self.set_text_color(0, 0, 0)  # Reset text color
        return query_text

        # def Section6_1 (con, output_path, filename, drug):

    def Section6_1(self, con, output_path, filename, drug):
        query_text = f"\n The Medical Dictionary for Regulatory Activities (MedDRA) version 26.1 is the coding dictionary utilised for the presentation of AE/Adverse Drug Reactions (ADR).\n The tabulation report (organised by System Organ Class and MedDRA PT) summarises each AE coincident with product rather than each report. This report is produced from a dynamic database, which changes over time as reports are updated, and reflects the most current data available at the time that it was generated. As a single report may contain both serious and non-serious and/or both listed and unlisted AEs, a report may be presented in more than one category under each source. Therefore, the sum of the total number of reports across sources may exceed the number of unique reports that exist overall. \n Attempts (defined as phone calls, letters, questionnaires) have been made by the company to request follow-up information and/or medical confirmation of AE reports. The data included within this report represent the most complete report information available at the time of analysis.\n"
        # Add the query text to the PDF
        self.set_font('Arial', '', 12)
        self.set_text_color(0, 0, 0)
        self.multi_cell(0, 10, query_text.encode('latin-1', errors='ignore').decode('latin-1'), 0,
                        1)  # Encode and decode using latin-1
        self.set_text_color(0, 0, 0)  # Reset text color
        return query_text

        # def Section6_2 (con, output_path, filename, drug):

    def Section6_2(self, con, output_path, filename, drug):
        query_text = f"\n A cumulative summary tabulation of serious adverse events (SAEs) from Company-sponsored CTs is provided in Appendix 3.\n"
        self.set_font('Arial', '', 12)
        self.set_text_color(0, 0, 0)
        self.multi_cell(0, 10, query_text.encode('latin-1', errors='ignore').decode('latin-1'), 0,
                        1)  # Encode and decode using latin-1
        self.set_text_color(0, 0, 0)  # Reset text color
        return query_text

        # def Section6_3 (con, output_path, filename, drug):

    def Section6_3(self, con, output_path, filename, drug):
        query_text = f"A cumulative (IBD to DLP) and interval {self.formatted_begin_date} to {self.formatted_end_date} summary tabulation of ADRs (serious and non-serious) is provided in Appendix 4. The serious and non-serious ADRs presented in this tabulation were derived from spontaneous individual case safety reports (ICSRs; from HCP, consumers, scientific literature, and RA) as well as serious ADRs from non-interventional studies.\n"
        self.set_font('Arial', '', 12)
        self.set_text_color(0, 0, 0)
        self.multi_cell(0, 10, query_text.encode('latin-1', errors='ignore').decode('latin-1'), 0,
                        1)  # Encode and decode using latin-1
        self.set_text_color(0, 0, 0)  # Reset text color
        return query_text

    def Section11(self, con, output_path, filename, drug):
      try:
        # def Section2(self, con, output_path, filename, drug):
        pubmed = PubMed(tool="PubMedSearcher", email="drruchijakhmola@gmail.com")
        drug = 'Allopurinol'
        # List of Keyword options
        keyword_options = [
            "pharmacovigilance",
            "adverse effect",
            "adverse event",
            "adverse reaction",
            "side effect",
            "contraindication",
            "interaction",
            "mutagenicity",
            "teratogenicity",
            "embriopath",
            "embryopath",
            "malformation",
            "pregnancy",
            "breastfeeding",
            "pediatric",
            "abuse",
            "misuse",
            "safety",
            "overdose",
            "overdosage",
            "poisoning",
            "suicide",
            "fatal",
            "risk",
            "dependence",
            "withdrawl",
            "accident",
            "medication error",
            "drug addiction",
            "disease aggravation",
            "disease progression",
            "lack of drug effect",
            "carcinogenicity",
            "toxic",
            "toxicity",
            "lactation",
            "intoxication",
            "maladministration"
        ]
        def check_for_adr(text):
            text = text.lower()  # Convert text to lowercase to make the search case-insensitive
            return any(keyword in text for keyword in adr_keywords)

        def get_sentiment(text):
            analysis = TextBlob(str(text))
            return analysis.sentiment.polarity

        # Define the categorization function
        def categorize_sentiment(score):
            if score < 0:
                return 'Negative'
            elif score > 0:
                return 'Positive'
            else:
                return 'Neutral'
        adr_keywords = ['adverse effect', 'side effect', 'drug reaction', 'complication', 'toxicity']
        # Initialize an empty list to store DataFrames
        all_dfs = []

        # Iterate over each Keyword option
        for keyword in keyword_options:
            # Construct search term
            search_term = f"({drug}) AND ({keyword})"

            # Query PubMed
            results = pubmed.query(search_term, max_results=9999)
            articleList = []
            articleInfo = []

            for article in results:
                articleDict = article.toDict()
                articleList.append(articleDict)

            for article in articleList:
                # time.sleep(1)
                pubmedId = article['pubmed_id'].partition('\n')[0]
                
                articleInfo.append({
                    'pubmed_id': pubmedId,
                    'title': article['title'],
                    'abstract': article['abstract'],
                    'publication_date': article['publication_date'],
                    'authors': article['authors']
                })

            # Generate Pandas DataFrame from list of dictionaries
            articlesPD = pd.DataFrame.from_dict(articleInfo)

            # Append current DataFrame to the list of DataFrames
            all_dfs.append(articlesPD)

        # Concatenate all DataFrames in the list into a single DataFrame
        all_articles_df = pd.concat(all_dfs, ignore_index=True)
        all_articles_df = all_articles_df.drop_duplicates(subset = 'pubmed_id')
        all_articles_df['publication_date'] = pd.to_datetime(all_articles_df['publication_date'])

        # Filter rows where 'publication_date' is between 'formatted_begin_date' and 'formatted_end_date'
        filtered_df = all_articles_df[(all_articles_df['publication_date'] >= self.formatted_begin_date) & (all_articles_df['publication_date'] <= self.formatted_end_date)]




        # Replace 'your_file_path/filtered_articles_PSURdate.csv' with the path to your CSV file
        file_path = 'https://raw.githubusercontent.com/Snigdhab26/Capstone/main/filtered_articles_PSURdate.csv'
        desired_filename = 'filtered_articles_PSURdate.csv'
        for f in filename:
            # Extract the filename from the full path
            file = os.path.basename(f)
            print(file)
            # Check if the extracted filename matches the desired filename
            if file == desired_filename:
                print("Found:", f)
                file_path = f
                break
            # else:
            #     print("File not found:", desired_filename)
        # Reading the CSV file into a DataFrame
        # filtered_df = pd.read_csv(file_path)
        filtered_df_cleaned = filtered_df.dropna(subset=['abstract'])

        # Assuming 'filtered_df_cleaned' is your DataFrame
        # Apply the sentiment analysis function
        filtered_df_cleaned['Sentiment Score'] = filtered_df_cleaned['abstract'].apply(get_sentiment)

        # Apply the categorization function
        filtered_df_cleaned['Sentiment Category'] = filtered_df_cleaned['Sentiment Score'].apply(categorize_sentiment)

        num_articles = len(filtered_df_cleaned)

        # Select and display only the desired columns
        df_subset = filtered_df_cleaned[['pubmed_id', 'title', 'abstract', 'Sentiment Score', 'Sentiment Category']]

        # List of keywords related to adverse drug reactions

        # Assuming 'filtered_df_cleaned' is your DataFrame
        df_subset['Contains ADR'] = df_subset['abstract'].apply(check_for_adr)

        # Display rows where an ADR is mentioned
        adr_present = df_subset[df_subset['Contains ADR'] == True]

        query_text = f"\n A literature search and analysis were performed utilizing MEDLINE® for abstracts covering the reporting interval {self.formatted_begin_date} to {self.formatted_end_date} to identify significant safety findings associated with {drug} . The search criteria used is mentioned in Appendix 8. During the reporting period, a total of {num_articles} abstracts were retrieved for the above-mentioned review period."
        self.set_font('Arial', '', 12)
        self.set_text_color(0, 0, 0)
        self.multi_cell(0, 10, query_text.encode('latin-1', errors='ignore').decode('latin-1'), 0,
                        1)  # Encode and decode using latin-1
        self.set_text_color(0, 0, 0)  # Reset text color
        return query_text

      except Exception as e:
          print(f"Section 17 : {e}")
          file_path = '/content/drive/MyDrive/ISB /Capstone/Data/Allopurinol/filtered_articles_PSURdate.csv'
          filtered_df = pd.read_csv(file_path)
          filtered_df_cleaned = filtered_df.dropna(subset=['abstract'])

          # Assuming 'filtered_df_cleaned' is your DataFrame
          # Apply the sentiment analysis function
          filtered_df_cleaned['Sentiment Score'] = filtered_df_cleaned['abstract'].apply(get_sentiment)

          # Apply the categorization function
          filtered_df_cleaned['Sentiment Category'] = filtered_df_cleaned['Sentiment Score'].apply(categorize_sentiment)

          num_articles = len(filtered_df_cleaned)

          # Select and display only the desired columns
          df_subset = filtered_df_cleaned[['pubmed_id', 'title', 'abstract', 'Sentiment Score', 'Sentiment Category']]

          # List of keywords related to adverse drug reactions

          # Assuming 'filtered_df_cleaned' is your DataFrame
          df_subset['Contains ADR'] = df_subset['abstract'].apply(check_for_adr)

          # Display rows where an ADR is mentioned
          adr_present = df_subset[df_subset['Contains ADR'] == True]

          query_text = f"\n A literature search and analysis were performed utilizing MEDLINE® for abstracts covering the reporting interval {self.formatted_begin_date} to {self.formatted_end_date} to identify significant safety findings associated with {drug} . The search criteria used is mentioned in Appendix 8. During the reporting period, a total of {num_articles} abstracts were retrieved for the above-mentioned review period."
          self.set_font('Arial', '', 12)
          self.set_text_color(0, 0, 0)
          self.multi_cell(0, 10, query_text.encode('latin-1', errors='ignore').decode('latin-1'), 0,
                          1)  # Encode and decode using latin-1
          self.set_text_color(0, 0, 0)  # Reset text color
          return query_text

    def Section17(self, con, output_path, filename, drug):
      try:
        # def Section2(self, con, output_path, filename, drug):
          drug ='Allopurinol'
        # List of Keyword options
          pubmed = PubMed(tool="PubMedSearcher", email="drruchijakhmola@gmail.com")

          # List of Keyword options
          keyword_options = [
              "efficacy",
              "effective",
              "ineffective",
              "intolerable",
              "not-tolerable",
              "discontinue",
              "non-effective",
              "satisfaction",
              "adherence",
              "compliance",
              "clinical-effect",
              "treatment-effect",
              "clinical benefit",
              "treatment effect",
              "clinical-benefit",
              "not tolerable",
              "non effective"
          ]

          # Initialize an empty list to store DataFrames
          all_dfs = []

          # Iterate over each Keyword option
          for keyword in keyword_options:
              # Construct search term
              search_term = f"({drug}) AND ({keyword})"

              # Query PubMed
              results = pubmed.query(search_term, max_results=9999)
              articleList = []
              articleInfo = []

              for article in results:
                  articleDict = article.toDict()
                  articleList.append(articleDict)

              for article in articleList:
                  pubmedId = article['pubmed_id'].partition('\n')[0]
                  articleInfo.append({
                      'pubmed_id': pubmedId,
                      'title': article['title'],
                      'abstract': article['abstract'],
                      'publication_date': article['publication_date'],
                      'authors': article['authors']
                  })

              # Generate Pandas DataFrame from list of dictionaries
              efficacy_articlesPD = pd.DataFrame.from_dict(articleInfo)

              # Append current DataFrame to the list of DataFrames
              all_dfs.append(efficacy_articlesPD)

          # Concatenate all DataFrames in the list into a single DataFrame
          all_articles_efficacy_df = pd.concat(all_dfs, ignore_index=True)
          all_articles_efficacy_df = all_articles_efficacy_df.drop_duplicates(subset = 'pubmed_id')
          # Export the combined DataFrame to a single CSV file
          export_efficacy_csv = all_articles_efficacy_df.to_csv(f'{drug}_all_keywords_pub.csv', index=None, header=True)


          # EDA on abstracts

          # 'all_articles_efficacy_df' is our DataFrame and 'formatted_begin_date' and 'formatted_end_date' are our date range variables for PSUR

          # First, ensure 'publication_date' is in datetime format
          all_articles_efficacy_df['publication_date'] = pd.to_datetime(all_articles_efficacy_df['publication_date'])

          # Filter rows where 'publication_date' is between 'formatted_begin_date' and 'formatted_end_date'
          filtered_efficacy_df = all_articles_efficacy_df[(all_articles_efficacy_df['publication_date'] >= self.formatted_begin_date) & (all_articles_efficacy_df['publication_date'] <= self.formatted_end_date)]


          # Remove rows where the 'abstract' is missing
          filtered_efficacy_df_cleaned = filtered_efficacy_df.dropna(subset=['abstract'])

          num_efficacy_articles = len(filtered_efficacy_df_cleaned)
          print ('num_efficacy_articles',num_efficacy_articles )

          # List of keywords related to efficacy
          efficacy_keywords = [
              "efficacy", "effective", "ineffective", "intolerable", "not-tolerable", "discontinue", "non-effective",
              "satisfaction", "adherence", "compliance", "clinical-effect", "treatment-effect", "clinical benefit",
              "treatment effect", "clinical-benefit", "not tolerable", "non effective"
          ]

          # Function to check if any keywords are in the abstract
          def check_for_efficacy(text):
              text = text.lower()  # Convert text to lowercase to make the search case-insensitive
              return any(keyword in text for keyword in efficacy_keywords)

          # Add a column indicating if the abstract mentions efficacy-related keywords
          filtered_efficacy_df_cleaned.loc[:, 'Is Efficacious'] = filtered_efficacy_df_cleaned['abstract'].apply(check_for_efficacy)

          # Filter for abstracts that mention efficacy-related keywords
          filtered_efficacy_df_cleaned = filtered_efficacy_df_cleaned[filtered_efficacy_df_cleaned['Is Efficacious'] == True]

          # Define a function to calculate sentiment scores using TextBlob

          def categorize_sentiment(score):
              if score < 0:
                  return 'Negative'
              elif score > 0:
                  return 'Positive'
              else:
                  return 'Neutral'  # or you can adjust the range for neutrality if needed



          # Define the sentiment analysis function
          def get_sentiment(text):
              analysis = TextBlob(str(text))
              return analysis.sentiment.polarity

          # Define the categorization function
          def categorize_sentiment(score):
              if score < 0:
                  return 'Negative'
              elif score > 0:
                  return 'Positive'
              else:
                  return 'Neutral'


          # Apply the sentiment analysis function
          filtered_efficacy_df_cleaned ['Sentiment Score'] = filtered_efficacy_df_cleaned ['abstract'].apply(get_sentiment)

          # Apply the categorization function
          filtered_efficacy_df_cleaned ['Sentiment Category'] = filtered_efficacy_df_cleaned ['Sentiment Score'].apply(categorize_sentiment)

          # Select and display only the desired columns
          df_subset = filtered_efficacy_df_cleaned [['pubmed_id', 'title', 'abstract', 'publication_date', 'Is Efficacious', 'Sentiment Category']]

          # Filter the DataFrame based on sentiment and efficacy conditions
          positive_efficacy_df = df_subset[
              (df_subset['Sentiment Category'] == 'Positive') &
              (df_subset['Is Efficacious'] == True)
          ]

          # save to a CSV file
          positive_efficacy_df.to_csv(f'{drug}_positive_efficacy_abstracts.csv', index=False)

          # PubMedBERT

          tokenizer = AutoTokenizer.from_pretrained("microsoft/BiomedNLP-PubMedBERT-base-uncased-abstract-fulltext")
          model = AutoModel.from_pretrained("microsoft/BiomedNLP-PubMedBERT-base-uncased-abstract-fulltext")



          def get_embeddings(abstracts):
              model.eval()  # Set the model to evaluation mode
              embeddings = []
              with torch.no_grad():  # Disable gradient calculations for efficiency
                  for abstract in abstracts:
                      inputs = tokenizer(abstract, return_tensors="pt", truncation=True, padding="max_length", max_length=512)
                      outputs = model(**inputs)
                      # Extract embeddings from the last hidden state (mean pooling here as an example)
                      mean_embeddings = outputs.last_hidden_state.mean(1)
                      embeddings.append(mean_embeddings)
              return embeddings

          # Select the top 10 latest abstracts by publication date
          # top_10_abstracts_df = filtered_efficacy_df_cleaned.sort_values(by='publication_date', ascending=False).head(10)
          # top_10_abstracts = top_10_abstracts_df['abstract'].tolist()

          top_10_abstracts_df1 = filtered_efficacy_df_cleaned.sort_values(by='publication_date', ascending=False).head(50)
          top_10_abstracts_df1 = top_10_abstracts_df1.drop_duplicates(subset=['abstract'])
          top_10_abstracts_df = top_10_abstracts_df1.sort_values(by='publication_date', ascending=False).head(10)
          top_10_abstracts = top_10_abstracts_df['abstract'].tolist()
          # Check if the top 10 abstracts have been correctly selected
          print("Number of top abstracts selected:", len(top_10_abstracts))

          # Check if the top 10 abstracts have been correctly selected
          print("Number of top abstracts selected:", len(top_10_abstracts))



          # Process only the top 10 abstracts
          embeddings = get_embeddings(top_10_abstracts)
          nltk.download('wordnet')
          print("Pipeline Started")
          # Load the summarization pipeline using BART
          summarizer = pipeline("summarization" , model="facebook/bart-large-cnn")
          print("Pipeline Ended")
          # Function to summarize an individual abstract
          def summarize_abstract(abstract, min_length=30, max_length=130):
              # Tokenize the abstract to count the words
              word_count = len(abstract.split())

              # Adjust max_length and min_length based on the actual word count of the abstract
              adjusted_max_length = min(max_length, word_count)
              adjusted_min_length = min(min_length, word_count // 2)

              # Handle cases where the abstract is too short to summarize
              if adjusted_max_length < 1:
                  return abstract
              print(abstract)
              # Generate summary using the adjusted lengths
              summary = summarizer(abstract, max_length=adjusted_max_length, min_length=adjusted_min_length, do_sample=False)
              return summary[0]['summary_text']
          print("Summarisation Started")
          # Generate summaries for the top 10 abstracts
          top_10_summaries = [summarize_abstract(abstract) for abstract in top_10_abstracts]
          query_text =f"\n \n  Information on both efficacy and effectiveness of the product at the beginning of the reporting period is provided in the following. This information forms the basis for the benefit evaluation.Indications of allopurinol are presented in Section 1.\n \n The following publications were received regarding the efficacy of {drug}:\n"

          for i, summary in enumerate(top_10_summaries):
            query_text = query_text + '\n' + 'Study {} Summary \n : {}'.format(i+1,summary)
          
          print("Summarisation Done")
          # Ensure you have downloaded the necessary NLTK data files
          nltk.download('punkt')
          # top_10_abstracts contains original abstracts and top_10_summaries contains generated summaries

          # Calculate BLEU score
          def calculate_bleu(reference, hypothesis):
              smoothie = SmoothingFunction().method4
              return sentence_bleu([reference.split()], hypothesis.split(), smoothing_function=smoothie)

          # Calculate ROUGE score
          def calculate_rouge(reference, hypothesis):
              scorer = rouge_scorer.RougeScorer(['rouge1', 'rouge2', 'rougeL'], use_stemmer=True)
              scores = scorer.score(reference, hypothesis)
              return scores

          # Calculate BERTScore
          def calculate_bertscore(references, hypotheses):
              P, R, F1 = bert_score(hypotheses, references, lang='en', rescale_with_baseline=True)
              return P.mean().item(), R.mean().item(), F1.mean().item()

          # Initialize BLEU from sacrebleu
          bleu = BLEU()

          # Tokenize text
          def tokenize(text):
              return nltk.word_tokenize(text)
          
          def evaluate_summarization(metrics_df):
              feedback = []

              for i, row in metrics_df.iterrows():
                  summary_quality = []
                  
                  # Check BLEU score
                  if row['BLEU'] < 0.1:
                      summary_quality.append("The summary has low exact match with the original text.")
                  elif row['BLEU'] < 0.3:
                      summary_quality.append("The summary has moderate exact match with the original text.")
                  else:
                      summary_quality.append("The summary has good exact match with the original text.")

                  # Check ROUGE-1 score
                  if row['ROUGE-1'] < 0.2:
                      summary_quality.append("The summary has low overlap with the original text.")
                  elif row['ROUGE-1'] < 0.5:
                      summary_quality.append("The summary has moderate overlap with the original text.")
                  else:
                      summary_quality.append("The summary has good overlap with the original text.")

                  # Check BERTScore F1
                  if row['BERTScore_F1'] < 0.2:
                      summary_quality.append("The summary is semantically quite different from the original text.")
                  elif row['BERTScore_F1'] < 0.5:
                      summary_quality.append("The summary captures some of the original meaning.")
                  else:
                      summary_quality.append("The summary captures most of the original meaning.")

                  # Check METEOR score
                  if row['METEOR'] < 0.1:
                      summary_quality.append("The summary has poor linguistic quality.")
                  elif row['METEOR'] < 0.3:
                      summary_quality.append("The summary has moderate linguistic quality.")
                  else:
                      summary_quality.append("The summary has good linguistic quality.")

                  # Combine feedback for this summary
                  feedback.append(f"Summary {i + 1} feedback:\n" + "\n".join(summary_quality) + "\n")
              
              return "\n".join(feedback)

          # Assuming metrics_df is already defined


          # Calculate metrics for each summary
          metrics = {
              'abstract': [],
              'summary': [],
              'BLEU': [],
              'ROUGE-1': [],
              'ROUGE-2': [],
              'ROUGE-L': [],
              'BERTScore_P': [],
              'BERTScore_R': [],
              'BERTScore_F1': [],
              'METEOR': []
          }

          for abstract, summary in zip(top_10_abstracts, top_10_summaries):
              # Tokenize abstract and summary
              tokenized_abstract = tokenize(abstract)
              tokenized_summary = tokenize(summary)

              # Calculate BLEU
              bleu_score = calculate_bleu(abstract, summary)

              # Calculate ROUGE
              rouge_scores = calculate_rouge(abstract, summary)

              # Calculate BERTScore
              P, R, F1 = calculate_bertscore([abstract], [summary])

              # Calculate METEOR using nltk
              meteor_score_value = nltk_meteor_score.meteor_score([tokenized_abstract], tokenized_summary)

              # Store metrics
              metrics['abstract'].append(abstract)
              metrics['summary'].append(summary)
              metrics['BLEU'].append(bleu_score)
              metrics['ROUGE-1'].append(rouge_scores['rouge1'].fmeasure)
              metrics['ROUGE-2'].append(rouge_scores['rouge2'].fmeasure)
              metrics['ROUGE-L'].append(rouge_scores['rougeL'].fmeasure)
              metrics['BERTScore_P'].append(P)
              metrics['BERTScore_R'].append(R)
              metrics['BERTScore_F1'].append(F1)
              metrics['METEOR'].append(meteor_score_value)

          # Create a DataFrame to store the metrics
          print("PDF PASTING")
          metrics_df = pd.DataFrame(metrics)
          summary_feedback = evaluate_summarization(metrics_df)
          self.set_font('Arial', '', 12)
          self.set_text_color(0, 0, 0)
          self.multi_cell(0, 10, query_text.encode('latin-1', errors='ignore').decode('latin-1'), 0,
                          1)  # Encode and decode using latin-1
          self.set_text_color(0, 0, 0)  # Reset text color

      except Exception as e:
          print(f"Section 17 : {e}")
          streamlit.write ("Scrapping limit exceeded")
          query_text = f"\n \n  Information on both efficacy and effectiveness of the product at the beginning of the reporting period is provided in the following. This information forms the basis for the benefit evaluation.Indications of {drug} are presented in Section 1.\n \n The following publications were received regarding the efficacy of DUMMYDRUG:\n"
          query_text = query_text 
          metrics_df = pd.DataFrame()

      return query_text , metrics_df,summary_feedback

    def extract_section(self, doc_path, start_text, end_text):
        document = Document(doc_path)
        extract = False
        extracted_lines = []

        # Remove tabs from the section texts
        start_text = start_text.replace('\t', '')
        end_text = end_text.replace('\t', '')

        # Iterate through each paragraph in the document
        for paragraph in document.paragraphs:
            # Normalize paragraph text by removing tabs
            paragraph_text = paragraph.text.replace('\t', '')

            # Print debugging information for each paragraph
            print(f"DEBUG: Paragraph text: {paragraph_text}")

            # Start extracting lines once the starting section is found
            if start_text in paragraph_text:
                print(f"DEBUG: Found start text '{start_text}'")
                extract = True
                continue
            # Stop extracting lines once the ending section is found
            if end_text in paragraph_text and extract:
                print(f"DEBUG: Found end text '{end_text}'")
                break
            # Copy the lines if extraction is enabled
            if extract:
                extracted_lines.append(paragraph_text)

        return "\n".join(extracted_lines)

    def extract_next_line(self, doc_path, marker_text):
        # Extract the line immediately following the marker text in a DOCX document."""
        # Load the DOCX document
        document = Document(doc_path)
        found_marker = False
        next_line = ""

        # Normalize marker text (remove tabs and extra spaces)
        marker_text = marker_text.replace('\t', ' ').strip()

        # Iterate through each paragraph in the document
        for paragraph in document.paragraphs:
            # Normalize paragraph text
            paragraph_text = paragraph.text.replace('\t', ' ').strip()

            # Start extracting the next line after finding the marker
            if found_marker:
                next_line = paragraph_text
                break

            if marker_text.lower() in paragraph_text.lower():
                print(f"DEBUG: Found marker text '{marker_text}'")
                found_marker = True

        return next_line

        # def Section1(self,con, content, output_path, filename, drug):
        #     try:

        #         _filepath = pdf.filename_extract(con, filename)
        #         print(_filepath)

    def parse_date(self, date_str):
        # Parse a date string into a valid datetime object using dateparser."""
        date = dateparser.parse(date_str, settings={'DATE_ORDER': 'DMY'})
        if not date:
            raise ValueError(f"Invalid date format: {date_str}")
        return date

    def add_table_to_docx(self, doc, df):
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Add the header rows
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = str(column)

        # Add the data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)

    def create_docx(self,content ,head, df = None):
        doc = Document()
        doc.add_heading(head.upper(), level=1)
        doc.add_paragraph(content)
        if df is not None:
            self.add_table_to_docx(doc, df)
        # Save the document to a BytesIO object
        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        
        return byte_io
    


    def Executive_Summary(self, con, output_path, filename, drug, country, IBD):
        try:
            # Path to your DOCX file
            doc_path = self.filename_extract('CDS', filename)
            print(doc_path)
            # Define the starting and ending section texts
            start_text = "Therapeutic indications"
            end_text = "Posology and method of administration"

            # Extract the desired section
            section_text = self.extract_section(doc_path, start_text, end_text)

            # Display the extracted section
            print("Extracted Section:")
            print(section_text)

            # Optionally, save the extracted section to a text file
            with open("extracted_section.txt", "w", encoding="utf-8") as file:
                file.write(section_text)
            # Define the marker text
            marker_text = "QUALITATIVE AND QUANTITATIVE COMPOSITION"

            # Extract the next line after the marker
            next_line_text = self.extract_next_line(doc_path, marker_text)

            marketed_count = 0
            approved_count = 0
            qt=''
            marketed_count, approved_count , qt = (self.Section2('WWMA', output_path, filename, drug))

            country = country  # input("Enter the country name (e.g., 'United States', 'India', etc.): ").strip().title()
            birth_date_str = IBD  # input("Enter the International Birth Date (e.g., '01-12-2024', '1st Jan 2022', etc.): ")

            # Parse the input date using the parse_date function
            international_birth_date = self.parse_date(birth_date_str)
            query_text = f" \n This is a Periodic Safety Update Report for {drug} (International Birth Date {international_birth_date.strftime('%d-%m-%Y')}) based on the first approval in {country}. This report follows Periodic Benefit Risk Evaluation Report ICH-E2C (R2) guidelines and Good Pharmacovigilance Practice guideline Module VII.\n It summarizes the safety data received and processed by Sandoz from worldwide sources for the period covering {self.formatted_begin_date} to {self.formatted_end_date}.\n This report includes data obtained from Sandoz and its license partner Novartis.\n {drug} is indicated for the treatment of following indications:\n {section_text} \n {next_line_text} \n {drug} is currently authorized in {approved_count} countries and marketed in {marketed_count} of these countries as of data lock point."
            # print(query_text)
            self.set_font('Arial', '', 12)
            self.set_text_color(0, 0, 0)
            self.multi_cell(0, 10, query_text.encode('latin-1', errors='ignore').decode('latin-1'), 0,
                            1)  # Encode and decode using latin-1
            self.set_text_color(0, 0, 0)  # Reset text color
        except ValueError as e:
            print(e)
        return query_text

def get_last_directory_files(directory_path):
    # # Get all directories and files in the given path
    # contents = os.listdir(directory_path)
    # # Filter out directories from the list
    # subdirectories = [item for item in contents if os.path.isdir(os.path.join(directory_path, item))]
    #
    # if len(subdirectories) == 0:
    #     # If no subdirectories found, list all files in the current directory
    #     files = [file for file in contents if os.path.isfile(os.path.join(directory_path, file))]
    #     return directory_path, files
    # else:
    #     # Recursively find the last directory
    #     last_directory = subdirectories[-1]
    #     last_directory_path = os.path.join(directory_path, last_directory)
    #     return get_last_directory_files(last_directory_path)
    # Get all directories and files in the given path
    contents = os.listdir(directory_path)

    # Filter out directories from the list, excluding .ipynb_checkpoints if present
    subdirectories = [item for item in contents if
                      os.path.isdir(os.path.join(directory_path, item)) and item != '.ipynb_checkpoints']

    if len(subdirectories) == 0:
        # If no subdirectories found (or all are .ipynb_checkpoints), list all files in the current directory
        files = [file for file in contents if os.path.isfile(os.path.join(directory_path, file))]
        return directory_path, files
    else:
        # Recursively find the last directory
        last_directory = subdirectories[-1]
        last_directory_path = os.path.join(directory_path, last_directory)
        return get_last_directory_files(last_directory_path)

def create_pdf(product_name,filename,begin_date_str,end_date_str):
    # streamlit.write(streamlit.session_state)
    # streamlit.session_state['Start'] = True
    # streamlit.write(streamlit.session_state)
    # streamlit.session_state['Start'] = True
    #directory_path = "/content/drive/MyDrive/ISB /Capstone/Data"
    streamlit.write("Select Product for PSUR: ", product_name)
    # directory_path = os.getcwd()
    # directory_path = os.path.join(directory_path, 'Drugs')
    # #streamlit.write(directory_path)
    # #directory_path = '/Users/harishbalajim/Documents/ISB AMPBA/Capstone Project/Drugs'
    # Products = [folder for folder in os.listdir(directory_path)]
    # #product_name = input('Please enter Product Name: ')
    # # directory_path = os.path.join(directory_path,product_name.title())
    # #Products = [folder for folder in os.listdir(directory_path)]
    # #streamlit.write("Select Product for PSUR: ", Products)
    # # streamlit.write('directory_path', directory_path)
    # # streamlit.write(product_name)
    # folders = [folder for folder in os.listdir(directory_path) if folder == product_name]
    # streamlit.write(folders)
    #streamlit.write('folders',folders)
    filename = ['https://raw.githubusercontent.com/Snigdhab26/Capstone/main/Drug/DUMMYDRUG/DUMMYDRUG_2_WWMA.xlsx','https://raw.githubusercontent.com/Snigdhab26/Capstone/main/Drug/DUMMYDRUG/DUMMYDRUG_3_Actions taken in the reporting interval for safety reasons.docx','https://raw.githubusercontent.com/Snigdhab26/Capstone/main/Drug/DUMMYDRUG/DUMMYDRUG_4Changes to reference safety information.docx','https://raw.githubusercontent.com/Snigdhab26/Capstone/main/Drug/DUMMYDRUG/DUMMYDRUG_Core Data Sheet_CDS.docx','https://raw.githubusercontent.com/Snigdhab26/Capstone/main/Drug/DUMMYDRUG/DUMMYDRUG_RMP.pdf']
    # if len(folders) > 0:
    #     target_dir = os.path.join(directory_path, folders[0])
    #     last_directory_path, last_directory_files = get_last_directory_files(target_dir)

    #     if last_directory_files:
    #         # print("Files in the last directory:", last_directory_path)
    #         for file in last_directory_files:
    #             # print(file)
    #             filename.append(os.path.join(last_directory_path, file))
    #             #streamlit.write('FILENAME',filename)
    #     else:
    #         streamlit.write("No files found in the last directory.")
    # else:
    #     streamlit.write("Product directory not found.")
    
    streamlit.write("Files are : \n ")
    for i, f in enumerate(filename,1):
       file_name_only = os.path.basename(f)
       print(streamlit.write(f"{i}. {file_name_only}"))
    # df = pd.read_excel("/content/Excel.xlsx")
    df2 = pd.read_excel("https://raw.githubusercontent.com/Snigdhab26/Capstone/main/Excel.xlsx",sheet_name='Sheet2')  # ("/content/Excel.xlsx")

    df1 = pd.read_excel("https://raw.githubusercontent.com/Snigdhab26/Capstone/main/Excel.xlsx", sheet_name='Sheet1')
    df1 = df1.fillna('')

    marketed = 0
    approved = 0
    country = 'US'
    birth_date_str = '01-01-2024'
    # Ask the user for the begin date
    #begin_date_str = '24-12-2020'
    # Ask the user for the end date
    #end_date_str = '23-12-2023'
    # Parse the begin date into a datetime object
    begin_date = pd.to_datetime(begin_date_str, errors='coerce')
    # Parse the end date into a datetime object
    end_date = pd.to_datetime(end_date_str, errors='coerce')
    # Check if parsing was successful
    if pd.isnull(begin_date) or pd.isnull(end_date):
        print("Invalid date format entered.")
    else:
        # Print the parsed begin and end dates in the specified format ("dd-Mon-yyyy")
        formatted_begin_date = begin_date.strftime("%d-%b-%Y")
        formatted_end_date = end_date.strftime("%d-%b-%Y")

    # Create a PDF object
    pdf = PDF(product_name,df1,df2,formatted_begin_date,formatted_end_date)
    # Set auto page break
    pdf.set_auto_page_break(auto=True, margin=15)
    # Output path for the PDF file
    output_path = f'pbrer_report_{product_name}.pdf'
    pdf.product_name = product_name
    # Add Page
    pdf.add_page()
    pdf.set_font('Arial', 'B', 8)
    pdf.cell(0, 8, 'TABLE OF CONTENTS'.encode('latin-1', 'replace').decode('latin-1'), 0, 1,'L')

    for index, row in df2.iterrows():

        link = pdf.add_link()
        #print(link)
        df2.at[index, 'Link'] = link
        if pd.isna(row['Content']):  # Check if 'Content' is NaN
            if pd.isna(row['Stage1']):
                content = row['Stage']
            else:
                content = row['Stage1']
        else:
            content = row['Content']

        sr_no = str(row['Sr No.']) + '.'  # Convert Sr No. to string

        if sr_no == '0.':
            pdf.Index_page(' ' + content,'L',0 , link,color=(0, 0, 255))
        elif sr_no.count('.') == 1:
            pdf.Index_page(sr_no + ' ' + content,'L',0 , link,color=(0, 0, 255))
        elif sr_no.count('.') == 2:  # Use chapter_body if sr_no has 2 or more '.'

            pdf.Index_page('       '+sr_no + ' ' + content,'L',20, link,color=(0, 0, 255))
            pdf.ln(1)
        else:

            pdf.Index_page('            '+ sr_no + ' ' + content,'L',8, link,color=(0, 0, 255))
            pdf.ln(1)


    pdf.set_text_color(0, 0, 0)

    for index, row in df2.iterrows():
        if pd.isna(row['Content']):  # Check if 'Content' is NaN
            if pd.isna(row['Stage1']):
                content = row['Stage']
            else:
                content = row['Stage1']
        else:
            content = row['Content']

        sr_no = str(row['Sr No.']) + '.'  # Convert Sr No. to string
        print(index,content)
        if sr_no == '0.':
                pdf.add_page()
                pdf.chapter_title( ' ' + content, row['Link'])
        elif  sr_no.count('.') == 1:
                pdf.add_page()
                pdf.chapter_title(sr_no + ' ' + content, row['Link'])
        else:  # sr_no.count('.') >= 2:  # Use chapter_body if sr_no has 2 or more '.'
                pdf.chapter_body(sr_no + ' ' + content, row['Link'])

        pdf.set_text_color(0, 0, 0)
        # content = 'INTRODUCTION'

        if 'Introduction'.upper() in content.upper():
            print(content.upper())
            streamlit.write(content.upper())
            Section1_text = pdf.Section1('CDS', content, output_path, filename, product_name)
            docx_file = pdf.create_docx(Section1_text,content)
            # streamlit.download_button(
            #           label="Download {}".format(content),
            #           data=docx_file,
            #           file_name="{}.docx".format(content),
            #           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            #       )
            streamlit.download_button(
                        label="Download {}".format(content),
                        data=docx_file,
                        file_name="{}.docx".format(content),  # Use a safe part of the content for the filename
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="{}_button_1".format(content)
                    )




        if 'ABBREVIATIONS' in content:
            # print(filename)
            # print(product_name)
            # print(output_path)
            print(content.upper())
            streamlit.write(content.upper())
            DF = pdf.Abbreviations('ABBREVIATIONS', output_path, filename)
            docx_file = pdf.create_docx('',content.upper(),df=DF)
            if streamlit.download_button(
            label="Download {}".format(content.upper()),
            data=docx_file,
            file_name="{}.docx".format(content),  # Use a safe part of the content for the filename
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="{}_button".format(content)
              ):
                  streamlit.session_state['GenerateButton'] = True
                  streamlit.session_state['Start'] = True
                  streamlit.session_state['Section1'] = True
                  streamlit.session_state['Exec_Summary'] = True

        if 'WORLDWIDE' in content:
            print(content)
            streamlit.write(content.upper())
            marketed, approved,Section2_text = pdf.Section2('WWMA', output_path, filename, product_name)
            docx_file = pdf.create_docx(Section2_text,content)
            if streamlit.download_button(
            label="Download {}".format(content.upper()),
            data=docx_file,
            file_name="{}.docx".format(content),  # Use a safe part of the content for the filename
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="{}_button".format(content)
              ):
                  streamlit.session_state['GenerateButton'] = True
                  streamlit.session_state['Start'] = True
                  streamlit.session_state['Section1'] = True
                  streamlit.session_state['Exec_Summary'] = True


        if 'Actions'.upper() in content.upper() and 'Reporting'.upper() in content.upper() and 'Reasons'.upper() in content.upper():
            print(content)
            streamlit.write(content.upper())
            Section3_text = pdf.Section3('Actions', content, output_path, filename, product_name)
            docx_file = pdf.create_docx(Section3_text,content)
            if streamlit.download_button(
            label="Download {}".format(content.upper()),
            data=docx_file,
            file_name="{}.docx".format(content),  # Use a safe part of the content for the filename
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="{}_button".format(content)
              ):
                  streamlit.session_state['GenerateButton'] = True
                  streamlit.session_state['Start'] = True
                  streamlit.session_state['Section1'] = True
                  streamlit.session_state['Exec_Summary'] = True

        if 'Changes'.upper() in content.upper():
            print(content)
            streamlit.write(content.upper())
            Section4_text = pdf.Section4('Changes', content, output_path, filename, product_name)
            docx_file = pdf.create_docx(Section4_text,content)
            if streamlit.download_button(
            label="Download {}".format(content.upper()),
            data=docx_file,
            file_name="{}.docx".format(content),  # Use a safe part of the content for the filename
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="{}_button".format(content)
              ):
                  streamlit.session_state['GenerateButton'] = True
                  streamlit.session_state['Start'] = True
                  streamlit.session_state['Section1'] = True
                  streamlit.session_state['Exec_Summary'] = True

        if 'DATA IN'.upper() in content.upper():
            print(content)
            streamlit.write(content.upper())
            Section6_text =pdf.Section6(content, output_path, filename, product_name)
            docx_file = pdf.create_docx(Section6_text,content)
            if streamlit.download_button(
            label="Download {}".format(content.upper()),
            data=docx_file,
            file_name="{}.docx".format(content),  # Use a safe part of the content for the filename
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="{}_button".format(content)
              ):
                  streamlit.session_state['GenerateButton'] = True
                  streamlit.session_state['Start'] = True
                  streamlit.session_state['Section1'] = True
                  streamlit.session_state['Exec_Summary'] = True

        if 'Reference Information'.upper() in content.upper():
            print(content)
            streamlit.write(content.upper())
            Section6_1_text = pdf.Section6_1(content, output_path, filename, product_name)
            docx_file = pdf.create_docx(Section6_1_text,content)
            if streamlit.download_button(
            label="Download {}".format(content.upper()),
            data=docx_file,
            file_name="{}.docx".format(content),  # Use a safe part of the content for the filename
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="{}_button".format(content)
              ):
                  streamlit.session_state['GenerateButton'] = True
                  streamlit.session_state['Start'] = True
                  streamlit.session_state['Section1'] = True
                  streamlit.session_state['Exec_Summary'] = True

        if 'Cumulative'.upper() in content.upper() and 'Adverse'.upper() in content.upper():
            print(content)
            streamlit.write(content.upper())
            Section6_2_text = pdf.Section6_2(content, output_path, filename, product_name)
            docx_file = pdf.create_docx(Section6_2_text,content)
            if streamlit.download_button(
            label="Download {}".format(content.upper()),
            data=docx_file,
            file_name="{}.docx".format(content),  # Use a safe part of the content for the filename
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="{}_button".format(content)
              ):
                  streamlit.session_state['GenerateButton'] = True
                  streamlit.session_state['Start'] = True
                  streamlit.session_state['Section1'] = True
                  streamlit.session_state['Exec_Summary'] = True


        if 'Cumulative'.upper() in content.upper() and 'Post-marketing'.upper() in content.upper():
            print(content)
            streamlit.write(content.upper())
            Section6_3_text = pdf.Section6_3(content, output_path, filename, product_name)
            docx_file = pdf.create_docx(Section6_3_text,content)
            if streamlit.download_button(
            label="Download {}".format(content.upper()),
            data=docx_file,
            file_name="{}.docx".format(content),  # Use a safe part of the content for the filename
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="{}_button".format(content)
              ):
                  streamlit.session_state['GenerateButton'] = True
                  streamlit.session_state['Start'] = True
                  streamlit.session_state['Section1'] = True
                  streamlit.session_state['Exec_Summary'] = True

        if 'Literature'.upper() in content.upper():
            # if 'Cumulative'.upper() in content.upper() and 'Post-marketing'.upper() in content.upper():
            print(content)
            streamlit.write(content.upper())
            Section11_text = pdf.Section11(content, output_path, filename, product_name)
            docx_file = pdf.create_docx(Section11_text,content)
            if streamlit.download_button(
            label="Download {}".format(content.upper()),
            data=docx_file,
            file_name="{}.docx".format(content),  # Use a safe part of the content for the filename
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="{}_button".format(content)
              ):
                  streamlit.session_state['GenerateButton'] = True
                  streamlit.session_state['Start'] = True
                  streamlit.session_state['Section1'] = True
                  streamlit.session_state['Exec_Summary'] = True

        if 'Baseline'.upper() in content.upper() and 'Efficacy'.upper() in content.upper():
            # if 'Cumulative'.upper() in content.upper() and 'Post-marketing'.upper() in content.upper():
            print(content)
            streamlit.write(content.upper())
            Section17_text,Section17_metrics , Summary = pdf.Section17(content, output_path, filename, product_name)
            docx_file = pdf.create_docx(Section17_text,content)
            print("PASTE")
            streamlit.download_button(
            label="Download {}".format(content.upper()),
            data=docx_file,
            file_name="{}.docx".format(content),  # Use a safe part of the content for the filename
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="{}_button".format(content)
              )

            if Section17_metrics.empty:
              print("Empty")
            else:
              print("Sucessfully done")
              streamlit.dataframe(Section17_metrics , height = 1200)
              streamlit.write(Summary)


        if 'Executive'.upper() in content.upper() and 'Summary'.upper() in content.upper():
            # if 'Cumulative'.upper() in content.upper() and 'Post-marketing'.upper() in content.upper():
            print(content)
            streamlit.write(content.upper())
            Exec_Summary = (pdf.Executive_Summary(content, output_path, filename, product_name, country, birth_date_str))
            docx_file = pdf.create_docx(Exec_Summary,content)
            if streamlit.download_button(
                        label="Download {}".format(content.upper()),
                        data=docx_file,
                        file_name="{}.docx".format(content),  # Use a safe part of the content for the filename
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="{}_button".format(content)
                    ):
                  streamlit.session_state['GenerateButton'] = True
                  streamlit.session_state['Start'] = True
                  streamlit.session_state['Section1'] = True
                  streamlit.session_state['Exec_Summary'] = True
                  
                  
    # Output the PDF
    pdf.output(output_path, 'F')
    #streamlit.write(output_path)
    with open(output_path, 'rb') as f:
        streamlit.download_button('Download Final PDF File', f, file_name=output_path)
        streamlit.session_state['DownloadPDF'] = True
